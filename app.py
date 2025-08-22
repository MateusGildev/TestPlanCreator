from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
import os

class CriadorPlanoTesteWord:
    def __init__(self):
        self.document = Document()
    
    def aplicar_fonte_arial(self, paragraph, tamanho_pt=11, negrito=False):
        """Aplica fonte Arial a todos os runs de um parágrafo"""
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(tamanho_pt)
            run.font.bold = negrito
    
    def coletar_multiplos_itens(self, campo_nome):
        """Coleta múltiplos itens para um campo específico"""
        print(f"\n--- {campo_nome.upper()} ---")
        print(f"Digite os {campo_nome.lower()} (digite 'fim' para terminar):")
        
        itens = []
        contador = 1
        
        while True:
            item = input(f"{campo_nome} {contador}: ").strip()
            if item.lower() == 'fim':
                break
            if item:
                itens.append(item)
                contador += 1
        
        return itens
    
    def coletar_dados(self):
        print("\n" + "="*60)
        print("CRIADOR DE PLANO DE TESTE - DOCUMENTO WORD")
        print("="*60)
        
        # Dados básicos
        id_plano = input("ID do Plano de Teste: ")
        nome = input("Nome do Plano de Teste: ")
        versao = input("Versão: ")
        
        # Validação de data
        while True:
            data_input = input("Data (dd/mm/yyyy): ")
            try:
                data_obj = datetime.strptime(data_input, "%d/%m/%Y")
                data_formatada = data_obj.strftime("%d/%m/%Y")
                break
            except ValueError:
                print("Formato de data inválido! Use dd/mm/yyyy")
        
        autor = input("Autor: ")
        
        # Coletar múltiplos itens
        objetivos = self.coletar_multiplos_itens("Objetivo")
        escopos = self.coletar_multiplos_itens("Escopo")
        estrategias = self.coletar_multiplos_itens("Estratégia de Teste")
        recursos = self.coletar_multiplos_itens("Recursos")
        
        return {
            'id_plano': id_plano,
            'nome': nome,
            'versao': versao,
            'data': data_formatada,
            'autor': autor,
            'objetivos': objetivos,
            'escopos': escopos,
            'estrategias': estrategias,
            'recursos': recursos
        }
    
    def criar_cabecalho(self, dados):
        """Cria o cabeçalho do documento"""
        # Título principal
        titulo = self.document.add_heading('PLANO DE TESTE', level=0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Configurar fonte do título manualmente
        self.aplicar_fonte_arial(titulo, tamanho_pt=16, negrito=True)
        
        # Adicionar espaço
        self.document.add_paragraph()
    
    def criar_tabela_informacoes(self, dados):
        """Cria tabela com informações básicas"""
        tabela = self.document.add_table(rows=5, cols=2)
        tabela.style = 'Table Grid'
        tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Dados da tabela
        info = [
            ('ID do Plano de Teste', dados['id_plano']),
            ('Nome do Plano', dados['nome']),
            ('Versão', dados['versao']),
            ('Data', dados['data']),
            ('Autor', dados['autor'])
        ]
        
        for i, (campo, valor) in enumerate(info):
            # Célula do campo
            cell_campo = tabela.cell(i, 0)
            cell_campo.text = campo
            # Configurar fonte para célula do campo
            for paragraph in cell_campo.paragraphs:
                self.aplicar_fonte_arial(paragraph, negrito=True)
            
            # Célula do valor
            cell_valor = tabela.cell(i, 1)
            cell_valor.text = str(valor)
            # Configurar fonte para célula do valor
            for paragraph in cell_valor.paragraphs:
                self.aplicar_fonte_arial(paragraph)
        
        # Adicionar espaço após a tabela
        self.document.add_paragraph()
    
    def criar_secao_com_lista(self, titulo, itens):
        """Cria uma seção com lista de itens"""
        if not itens:
            return
        
        # Título da seção
        heading = self.document.add_heading(titulo, level=1)
        # Configurar fonte do heading
        self.aplicar_fonte_arial(heading, tamanho_pt=12, negrito=True)
        
        # Adicionar itens da lista
        for item in itens:
            paragraph = self.document.add_paragraph(item, style='List Bullet')
            # Configurar fonte dos itens da lista
            self.aplicar_fonte_arial(paragraph)
        
        # Adicionar espaço
        self.document.add_paragraph()
    
    def criar_documento(self, dados):
        """Cria o documento Word completo"""
        # Configurar margens
        sections = self.document.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Criar conteúdo
        self.criar_cabecalho(dados)
        self.criar_tabela_informacoes(dados)
        self.criar_secao_com_lista("1. OBJETIVO", dados['objetivos'])
        self.criar_secao_com_lista("2. ESCOPO", dados['escopos'])
        self.criar_secao_com_lista("3. ESTRATÉGIA DE TESTE", dados['estrategias'])
        self.criar_secao_com_lista("4. RECURSOS", dados['recursos'])
        
        # Adicionar rodapé com data de geração
        footer = self.document.sections[0].footer
        paragraph = footer.paragraphs[0]
        paragraph.text = f"Documento gerado automaticamente em {datetime.now().strftime('%d/%m/%Y %H:%M')}"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Configurar fonte do rodapé
        self.aplicar_fonte_arial(paragraph, tamanho_pt=9)
    
    def salvar_documento(self, dados):
        """Salva o documento Word"""
        nome_arquivo = f"Plano_Teste_{dados['id_plano']}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        nome_arquivo = input(f"\nNome do arquivo Word (padrão: {nome_arquivo}): ") or nome_arquivo
        
        # Garantir extensão .docx
        if not nome_arquivo.lower().endswith('.docx'):
            nome_arquivo += '.docx'
        
        self.criar_documento(dados)
        self.document.save(nome_arquivo)
        
        return nome_arquivo
    
    def executar(self):
        print("=== CRIADOR DE PLANO DE TESTE - DOCUMENTO WORD ===")
        print("Para cada campo, você pode adicionar múltiplos itens.")
        print("Digite 'fim' para terminar a lista de cada campo.\n")
        
        while True:
            dados = self.coletar_dados()
            
            # Mostrar prévia
            print("\n" + "="*50)
            print("PRÉVIA DOS DADOS:")
            print("="*50)
            for key, value in dados.items():
                if isinstance(value, list):
                    print(f"{key.upper()}: {len(value)} item(s)")
                else:
                    print(f"{key.upper()}: {value}")
            
            # Confirmar
            confirmacao = input("\nDeseja gerar o documento com esses dados? (s/n): ").lower()
            if confirmacao == 's':
                nome_arquivo = self.salvar_documento(dados)
                caminho_absoluto = os.path.abspath(nome_arquivo)
                
                print(f"\n✅ DOCUMENTO CRIADO COM SUCESSO!")
                print(f"📁 Arquivo: {nome_arquivo}")
                print(f"📍 Local: {caminho_absoluto}")
                break
            else:
                print("Dados descartados. Vamos recomeçar...")
                self.document = Document()  # Resetar documento


def main():    
    try:
        while True:
            criador = CriadorPlanoTesteWord()
            criador.executar()
            
            continuar = input("\nDeseja criar outro plano? (s/n): ").lower()
            if continuar != 's':
                print("Programa finalizado. Obrigado!")
                break
                
    except KeyboardInterrupt:
        print("\n\nPrograma interrompido pelo usuário.")
    except Exception as e:
        print(f"❌ Ocorreu um erro: {e}")

if __name__ == "__main__":
    main()